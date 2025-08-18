from pathlib import Path
import pyperclip
from docx import Document
from docx.shared import Inches
from PIL import ImageGrab, Image
import io
import sys

OUT = Path(__file__).parent / "whatsapp_message.docx"

def save_clipboard_to_docx(out_path: Path):
    text = pyperclip.paste() or ""
    img = ImageGrab.grabclipboard()

    if not text and img is None:
        print("Clipboard is empty or contains unsupported data. Copy the WhatsApp message (text + image) and run again.")
        return

    doc = Document()

    if text:
        # keep original newlines as separate paragraphs
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.add_paragraph("")  # spacer

    # img can be a PIL Image, or a list of filenames, or None
    if isinstance(img, Image.Image):
        bio = io.BytesIO()
        img.save(bio, format="PNG")
        bio.seek(0)
        doc.add_picture(bio, width=Inches(4))  # adjust width as needed
    elif isinstance(img, list):
        # list of file paths (strings) copied from explorer
        for fp in img:
            try:
                doc.add_picture(fp, width=Inches(4))
            except Exception:
                pass

    doc.save(str(out_path))
    print(f"Wrote: {out_path}")

if __name__ == "__main__":
    save_clipboard_to_docx(OUT)